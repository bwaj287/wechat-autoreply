from __future__ import annotations

from pathlib import Path

from docx import Document

from erge_gateway.schemas import IngestResult


def ingest_docx(path: Path) -> IngestResult:
    doc = Document(str(path))
    text_chunks: list[str] = []
    headings = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = str(getattr(paragraph.style, "name", "") or "")
        if style_name.lower().startswith("heading"):
            headings += 1
        text_chunks.append(text)
    table_summaries: list[str] = []
    for idx, table in enumerate(doc.tables[:5], start=1):
        rows = []
        for row in table.rows[:5]:
            rows.append(" | ".join(cell.text.strip() for cell in row.cells))
        table_summaries.append(f"Table {idx}:\n" + "\n".join(rows))
    return IngestResult(
        kind="docx",
        source_path=str(path),
        text_chunks=text_chunks + table_summaries,
        metadata={
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "heading_count": headings,
            "extraction_mode": "text_native",
        },
        routing_hint="docx_structured",
    )
